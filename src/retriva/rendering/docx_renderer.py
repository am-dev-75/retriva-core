# Copyright (C) 2026 Andrea Marson (am.dev.75@gmail.com)
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#         http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

from pathlib import Path
from typing import Dict, Optional
from retriva.rendering import register_renderer
from retriva.logger import get_logger

logger = get_logger(__name__)

class DocxRenderer:
    """Renderer for Word (.docx) artifacts using python-docx."""

    def render(
        self,
        artifact_type: str,
        parameters: Dict[str, str],
        output_path: Path,
        cancel_check: Optional[callable] = None,
    ) -> bool:
        logger.info(f"Rendering docx artifact: {artifact_type}")
        
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx not installed. Cannot render docx.")
            return False

        title = parameters.get("title", "Retriva Artifact")
        content = parameters.get("content", "No content provided.")
        
        doc = Document()
        doc.add_heading(title, 0)
        doc.add_paragraph(content)
        
        if cancel_check and cancel_check():
            return False
            
        doc.save(str(output_path))
        return True

# Register the renderer
register_renderer("docx", DocxRenderer)
